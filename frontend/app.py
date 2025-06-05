import streamlit as st
import os
import sys
import time
import tempfile
from docx import Document

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from backend.rag_engine import RAGEngine
from backend.rfp_processor import RFPProcessor

import requests
from PIL import Image

USERS = {"admin": "password123", "user": "secret"}


# Login form UI with default Streamlit styling
def show_login_page():
    st.markdown(
        """
        <style>
            .main-title {
                font-size: 48px;
                font-weight: bold;
                text-align: center;
                color: inherit;
            }
            .login-container {
                background-color: #f9f9f9;
                padding: 10px;
                border-radius: 12px;
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
            }
            .login-header {
                font-size: 28px;
                font-weight: 600;
                color: inherit;
                margin-bottom: 10px;
            }
            .stTextInput>div>div>input {
                background-color: #ffffff;
            }
            .stButton button {
                font-weight: bold;
                border-radius: 8px;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )
    import base64
    from io import BytesIO

    # Load and convert the image to base64
    def get_base64_image(image_path):
        img = Image.open(image_path)
        buffered = BytesIO()
        img.save(buffered, format="PNG")  # or JPEG
        img_b64 = base64.b64encode(buffered.getvalue()).decode()
        return img_b64

    # Get base64 string
    img_b64 = get_base64_image("smartfill-logo.jpeg")

    # Display the image centered using HTML
    st.markdown(
        f"""
        <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 0px;">
            <img src="data:image/png;base64,{img_b64}" width="90"/>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="main-title">SMARTFILL AI</div>', unsafe_allow_html=True)
    st.divider()

    # Centered layout using columns
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        # st.markdown('<div class="login-container">', unsafe_allow_html=True)
        st.markdown('<div class="login-header">🔐 Login</div>', unsafe_allow_html=True)

        username = st.text_input("Username", placeholder="Enter your username")
        password = st.text_input(
            "Password", type="password", placeholder="Enter your password"
        )
        login_button = st.button("Login")

        if login_button:
            if not username or not password:
                st.warning("⚠️ Please enter both username and password.")
                return

            try:
                if USERS.get(username) == password:
                    st.session_state.logged_in = True
                    st.success("✅ Login successful!")
                    st.rerun()
                else:
                    st.error("❌ Invalid credentials.")
            except requests.exceptions.RequestException:
                st.error("⚠️ Failed to connect to authentication server.")
        st.markdown("</div>", unsafe_allow_html=True)


# If not logged in, show login page
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    show_login_page()
    st.stop()

st.set_page_config(page_title="SMARTFILL AI", layout="wide")

# Initialize session state
if "rag" not in st.session_state:
    st.session_state.rag = RAGEngine(layout_aware=True)
if "rfp_processor" not in st.session_state:
    st.session_state.rfp_processor = RFPProcessor()
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "indexed" not in st.session_state:
    st.session_state.indexed = False
if "needs_context" not in st.session_state:
    st.session_state.needs_context = []
if "answers" not in st.session_state:
    st.session_state.answers = {}
if "current_question_idx" not in st.session_state:
    st.session_state.current_question_idx = 0
if "questions_mode" not in st.session_state:
    st.session_state.questions_mode = None
if "export_status" not in st.session_state:
    st.session_state.export_status = {"ready": False, "filename": None, "error": None}
if "document_generated" not in st.session_state:
    st.session_state.document_generated = False

# Critical fix: Initialize proper email state variables
if "email_form_visible" not in st.session_state:
    st.session_state.email_form_visible = False
if "email_recipient" not in st.session_state:
    st.session_state.email_recipient = ""
if "current_doc_filename" not in st.session_state:
    st.session_state.current_doc_filename = None
if "email_success" not in st.session_state:
    st.session_state.email_success = False
if "email_success_message" not in st.session_state:
    st.session_state.email_success_message = ""


# Define callbacks for email functionality
def show_email_form():
    st.session_state.email_form_visible = True
    st.session_state.email_success = False  # Reset success state when showing form


def hide_email_form():
    st.session_state.email_form_visible = False


def set_email_success(message):
    st.session_state.email_success = True
    st.session_state.email_success_message = message
    # Hide the form after 3 seconds
    time.sleep(3)
    st.session_state.email_form_visible = False


def send_email_now(recipient, filename):
    if recipient.strip():
        try:
            from backend.utils import send_email

            success, message = send_email(
                recipient_email=recipient,
                subject="SMARTFILL AI - RFP Responses",
                body="Attached are the SMARTFILL AI generated RFP responses.",
                attachment_path=filename,
            )
            if success:
                success_msg = f"✅ Document successfully delivered to {recipient}!"
                set_email_success(success_msg)
            return success, message
        except Exception as e:
            return False, str(e)
    else:
        return False, "Please enter a valid email address"


rag = st.session_state.rag
rfp_processor = st.session_state.rfp_processor


# Define callback functions for mode selection
def set_all_mode():
    st.session_state.questions_mode = "all"
    st.session_state.current_question_idx = 0


def set_single_mode():
    st.session_state.questions_mode = "single"
    st.session_state.current_question_idx = 0


def clear_input_fields():
    # Clear any existing context input for the current question
    current_context_key = f"context_{st.session_state.current_question_idx}"
    if current_context_key in st.session_state:
        st.session_state[current_context_key] = ""

    # Clear any existing file upload for the current question
    current_file_key = f"file_{st.session_state.current_question_idx}"
    if current_file_key in st.session_state:
        st.session_state[current_file_key] = None


def generate_document(answers):
    """Generate a Word document with RFP responses"""
    try:
        # Create Word document
        doc = Document()
        doc.add_heading("SMARTFILL AI - RFP Responses", 0)
        doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Questions: {len(answers)}")
        doc.add_paragraph()

        # Add each Q&A to document
        for q, a in answers.items():
            doc.add_heading("Question", level=2)
            doc.add_paragraph(q)
            doc.add_heading("Response", level=2)
            doc.add_paragraph(a)
            doc.add_paragraph()

        # Save to file
        import uuid

        filename = f"RFP_Responses_{uuid.uuid4().hex[:8]}.docx"
        doc.save(filename)
        st.session_state.current_doc_filename = filename

        return filename, None
    except Exception as e:
        return None, str(e)


# Sidebar for document processing
with st.sidebar:
    st.title("📂 Upload Training Documents")

    training_files = st.file_uploader(
        "Upload your training documents",
        type=["pdf", "xlsx", "docx", "csv"],
        accept_multiple_files=True,
        help="Upload any PDF, Excel, Word, or CSV files for training",
    )

    if st.button("Process Documents"):
        if training_files:
            status = st.empty()
            progress_text = st.empty()
            progress_bar = st.progress(0)
            total_files = len(training_files)

            for i, file in enumerate(training_files):
                try:
                    progress_text.text(f"Processing: {file.name}")

                    # Handle Excel files specially for Q&A pairs if they match the expected format
                    if file.name.endswith(".xlsx"):
                        with tempfile.NamedTemporaryFile(
                            delete=False, suffix=".xlsx"
                        ) as tmp:
                            tmp.write(file.read())
                            # Try to process as Q&A pairs first
                            try:
                                rfp_processor.load_training_qa_pairs(tmp.name)
                                st.info(f"✅ Processed Q&A pairs from {file.name}")
                            except Exception as e:
                                # If Q&A processing fails, process as regular document
                                chunks_added = rag.add_document(file)
                                st.info(f"✅ Processed {file.name} as regular document")
                            os.remove(tmp.name)
                    else:
                        # Process all other document types
                        chunks_added = rag.add_document(file)
                        progress_text.text(
                            f"✅ Added {chunks_added} chunks from {file.name}"
                        )

                except Exception as e:
                    progress_text.text(f"❌ Error processing {file.name}: {str(e)}")
                    time.sleep(2)

                progress_bar.progress((i + 1) / total_files)
                time.sleep(0.5)

            progress_bar.empty()
            progress_text.empty()
            status.success("✅ All documents processed successfully")
            st.session_state.indexed = True
        else:
            st.warning("Please upload at least one document to process")

# Display success message if email was sent successfully
if st.session_state.email_success:
    st.success(st.session_state.email_success_message)
    # Reset the success flag after displaying the message
    time.sleep(3)
    st.session_state.email_success = False
    st.rerun()

# Main UI
st.title("🤖 SMARTFILL AI - RFP Assistant")

if st.session_state.indexed:
    st.markdown("### 📑 Upload RFP Document")
    rfp_file = st.file_uploader(
        "Please Upload Files here",
        type=["xlsx"],
        key="rfp",
        help="Upload your RFP Excel file here",
    )

    if rfp_file and st.button("Process RFP"):
        with st.spinner("Processing RFP..."):
            # Reset state for new RFP
            st.session_state.questions_mode = None
            st.session_state.current_question_idx = 0
            st.session_state.answers = {}
            st.session_state.document_generated = False
            st.session_state.email_form_visible = False

            # Process RFP
            answers, needs_context = rfp_processor.process_rfp(rfp_file, rag.db)
            st.session_state.answers = answers
            st.session_state.needs_context = needs_context

            # If there are no questions needing context, show success message
            if not needs_context:
                st.success("✅ All questions have been processed automatically!")
                st.session_state.show_final_button = True
            else:
                st.warning(f"⚠️ {len(needs_context)} questions need additional context")
                st.info(
                    "Please select how you would like to provide context for the questions:"
                )

    # Show Get Final Responses button if appropriate
    if "show_final_button" in st.session_state and st.session_state.show_final_button:
        if st.button("Get Final Responses", key="get_final_responses_auto"):
            try:
                with st.spinner("Generating final document..."):
                    filename, error = generate_document(st.session_state.answers)
                    if error:
                        st.error(f"❌ Failed to generate document: {error}")
                    else:
                        st.success(
                            "✅ Completed final question and responses generation!"
                        )

                        # Create two columns for export and email options
                        col1, col2 = st.columns(2)

                        with col1:
                            with open(filename, "rb") as file:
                                file_data = file.read()
                                st.download_button(
                                    label="📥 Export Responses as DOCX",
                                    data=file_data,
                                    file_name=filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                )

                        # The email button with on_click callback
                        with col2:
                            st.button(
                                "✉️ Send via Email",
                                key="email_button_1",
                                on_click=show_email_form,
                            )
            except Exception as e:
                st.error(f"❌ Failed to generate export: {str(e)}")

    # Question display mode selection
    if st.session_state.needs_context:
        col1, col2 = st.columns(2)
        with col1:
            st.checkbox(
                "Show all questions at once",
                key="show_all",
                on_change=set_all_mode,
                value=(st.session_state.questions_mode == "all"),
            )
        with col2:
            st.checkbox(
                "Show questions one by one",
                key="show_one",
                on_change=set_single_mode,
                value=(st.session_state.questions_mode == "single"),
            )

        # Handle mode selection conflicts
        if st.session_state.get("show_all") and st.session_state.get("show_one"):
            st.error("Please select only one display mode")
            st.session_state.questions_mode = None

        # Display questions based on selected mode
        if st.session_state.questions_mode:
            unanswered = [
                q
                for q in st.session_state.needs_context
                if q not in st.session_state.answers
            ]

            if unanswered:
                if st.session_state.questions_mode == "all":
                    st.markdown("### Questions Needing Additional Context")
                    question_containers = {}  # Store containers for each question

                    for question in unanswered:
                        question_containers[question] = st.container()
                        with question_containers[question]:
                            st.markdown(f"**Q:** {question}")
                            col1, col2, col3 = st.columns([2, 1, 1])
                            with col1:
                                context = st.text_area(
                                    "Provide additional context:",
                                    key=f"context_{question}",
                                )
                            with col2:
                                context_file = st.file_uploader(
                                    "Or upload a file with context",
                                    key=f"file_{question}",
                                    label_visibility="collapsed",
                                )
                            with col3:
                                # Add Update Context button for each question
                                if st.button(
                                    "Update Context", key=f"update_{question}"
                                ):
                                    if context.strip() or context_file:
                                        if context_file:
                                            chunks_added = rag.add_document(
                                                context_file
                                            )
                                            docs = rag.db.similarity_search(
                                                question, k=3
                                            )
                                            context = "\n\n".join(
                                                [doc.page_content for doc in docs]
                                            )

                                        response = rfp_processor.openai_client.chat.completions.create(
                                            model="gpt-4",
                                            messages=[
                                                {
                                                    "role": "system",
                                                    "content": "Generate a precise answer based on the provided context.",
                                                },
                                                {
                                                    "role": "user",
                                                    "content": f"Context:\n{context}\n\nQuestion: {question}",
                                                },
                                            ],
                                        )
                                        st.session_state.answers[question] = (
                                            response.choices[0].message.content
                                        )
                                        st.success("Answer updated successfully!")

                                        # Show preview of the answer
                                        st.markdown("**Generated Answer:**")
                                        st.markdown(st.session_state.answers[question])

                elif st.session_state.questions_mode == "single":
                    if st.session_state.current_question_idx < len(unanswered):
                        question = unanswered[st.session_state.current_question_idx]
                        st.markdown("### Current Question")
                        st.markdown(f"**Q:** {question}")

                        # Add question counter
                        st.markdown(
                            f"Question {st.session_state.current_question_idx + 1} of {len(unanswered)}"
                        )

                        # Generate unique keys for each question's input fields
                        context_key = f"context_{st.session_state.current_question_idx}"
                        file_key = f"file_{st.session_state.current_question_idx}"

                        col1, col2, col3 = st.columns([2, 1, 1])
                        with col1:
                            context = st.text_area(
                                "Provide additional context:", key=context_key
                            )
                        with col2:
                            context_file = st.file_uploader(
                                "Or upload a file with context",
                                key=file_key,
                                label_visibility="collapsed",
                            )
                        with col3:
                            # Add Update Context button
                            if st.button("Update Context"):
                                if context.strip() or context_file:
                                    if context_file:
                                        chunks_added = rag.add_document(context_file)
                                        docs = rag.db.similarity_search(question, k=3)
                                        context = "\n\n".join(
                                            [doc.page_content for doc in docs]
                                        )

                                    # Generate answer without context validation
                                    response = rfp_processor.openai_client.chat.completions.create(
                                        model="gpt-4",
                                        messages=[
                                            {
                                                "role": "system",
                                                "content": "Generate a precise answer based on the provided context.",
                                            },
                                            {
                                                "role": "user",
                                                "content": f"Context:\n{context}\n\nQuestion: {question}",
                                            },
                                        ],
                                    )
                                    st.session_state.answers[question] = (
                                        response.choices[0].message.content
                                    )
                                    st.success("Answer updated successfully!")

                                    # Show preview of the answer
                                    st.markdown("**Generated Answer:**")
                                    st.markdown(st.session_state.answers[question])

                        # Add navigation buttons in a new row
                        nav_col1, nav_col2, nav_col3 = st.columns([1, 1, 1])
                        with nav_col1:
                            if st.session_state.current_question_idx > 0:
                                if st.button("⬅️ Previous Question"):
                                    clear_input_fields()  # Clear current input fields
                                    st.session_state.current_question_idx -= 1
                                    st.rerun()
                        with nav_col3:
                            if (
                                st.session_state.current_question_idx
                                < len(unanswered) - 1
                            ):
                                if st.button("Next Question ➡️"):
                                    clear_input_fields()  # Clear current input fields
                                    st.session_state.current_question_idx += 1
                                    st.rerun()

                        # Only show the current question's answer if it exists
                        if question in st.session_state.answers:
                            st.markdown("**Current Answer:**")
                            st.markdown(st.session_state.answers[question])
                    else:
                        st.success("✅ All questions have been processed!")
            else:
                st.success("✅ All questions have been processed!")

            # Show Get Final Responses button only when appropriate
            if not st.session_state.needs_context or not unanswered:
                st.markdown("---")
                if st.button("Get Final Responses", key="get_final_responses"):
                    try:
                        with st.spinner("Generating final document..."):
                            filename, error = generate_document(
                                st.session_state.answers
                            )
                            if error:
                                st.error(f"❌ Failed to generate document: {error}")
                            else:
                                st.success(
                                    "✅ Completed final question and responses generation!"
                                )

                                # Create two columns for export and email options
                                col1, col2 = st.columns(2)

                                with col1:
                                    with open(filename, "rb") as file:
                                        file_data = file.read()
                                        st.download_button(
                                            label="📥 Export Responses as DOCX",
                                            data=file_data,
                                            file_name=filename,
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        )

                                # The email button uses the on_click parameter for the callback
                                with col2:
                                    st.button(
                                        "✉️ Send via Email",
                                        key="email_button_2",
                                        on_click=show_email_form,
                                    )
                    except Exception as e:
                        st.error(f"❌ Failed to generate export: {str(e)}")

# EMAIL FORM SECTION - This is a completely separate section at the end
# Only show if the email form should be visible
if st.session_state.email_form_visible and st.session_state.current_doc_filename:
    # Create a container with a border to make it stand out
    email_container = st.container()

    with email_container:
        st.markdown(
            """
        <style>
        .email-section {
            border: 1px solid #ddd;
            border-radius: 5px;
            padding: 15px;
            margin: 10px 0;
            background-color: #f8f9fa;
        }
        .success-message {
            background-color: #d4edda;
            color: #155724;
            padding: 10px;
            border-radius: 5px;
            margin: 10px 0;
            text-align: center;
            font-weight: bold;
            animation: fadeIn 0.5s;
        }
        @keyframes fadeIn {
            from { opacity: 0; }
            to { opacity: 1; }
        }
        </style>
        <div class="email-section">
        <h3>✉️ Send Document via Email</h3>
        </div>
        """,
            unsafe_allow_html=True,
        )

        # Email input field
        recipient_email = st.text_input(
            "Enter recipient's email address:", key="email_recipient_input"
        )

        # Create columns for send and cancel buttons
        button_col1, button_col2, button_col3 = st.columns([1, 1, 2])

        with button_col1:
            if st.button("📨 Send Email", key="send_email_btn"):
                if recipient_email:
                    with st.spinner("Sending email..."):
                        success, message = send_email_now(
                            recipient=recipient_email,
                            filename=st.session_state.current_doc_filename,
                        )

                        if success:
                            # Display a beautiful success message
                            st.markdown(
                                f"""
                            <div class="success-message">
                                <div>✅ Document successfully delivered to {recipient_email}!</div>
                                <div style="font-size: 0.9em; margin-top: 5px;">The recipient will receive the document shortly.</div>
                            </div>
                            """,
                                unsafe_allow_html=True,
                            )

                            # Set the success state and schedule the form to hide
                            st.session_state.email_success = True
                            st.session_state.email_success_message = f"✅ Document successfully delivered to {recipient_email}!"
                            st.session_state.email_form_visible = False
                            st.rerun()
                        else:
                            st.error(f"❌ Failed to send email: {message}")
                else:
                    st.warning("⚠️ Please enter a recipient email address")

        with button_col2:
            if st.button("❌ Cancel", key="cancel_email_btn"):
                st.session_state.email_form_visible = False
                st.rerun()

# import streamlit as st
# import pandas as pd
# import plotly.express as px
# import plotly.graph_objects as go
# import os
# import sys
# import time
# import tempfile
# import uuid
# from datetime import datetime, timedelta
# from docx import Document

# sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
# from backend.rag_engine import RAGEngine
# from backend.rfp_processor import RFPProcessor
# import requests
# from PIL import Image

# USERS = {
#     "admin": "password123",
#     "user": "secret"
# }

# # Login form UI with default Streamlit styling
# def show_login_page():
#     st.markdown(
#         """
#         <style>
#             .main-title {
#                 font-size: 48px;
#                 font-weight: bold;
#                 text-align: center;
#                 color: inherit;
#             }
#             .login-container {
#                 background-color: #f9f9f9;
#                 padding: 10px;
#                 border-radius: 12px;
#                 box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
#             }
#             .login-header {
#                 font-size: 28px;
#                 font-weight: 600;
#                 color: inherit;
#                 margin-bottom: 10px;
#             }
#             .stTextInput>div>div>input {
#                 background-color: #ffffff;
#             }
#             .stButton button {
#                 font-weight: bold;
#                 border-radius: 8px;
#             }
#         </style>
#         """,
#         unsafe_allow_html=True
#     )
#     import base64
#     from io import BytesIO

#     # Load and convert the image to base64
#     def get_base64_image(image_path):
#         img = Image.open(image_path)
#         buffered = BytesIO()
#         img.save(buffered, format="PNG")  # or JPEG
#         img_b64 = base64.b64encode(buffered.getvalue()).decode()
#         return img_b64

#     # Get base64 string
#     img_b64 = get_base64_image("smartfill-logo.jpeg")

#     # Display the image centered using HTML
#     st.markdown(
#         f"""
#         <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 0px;">
#             <img src="data:image/png;base64,{img_b64}" width="90"/>
#         </div>
#         """,
#         unsafe_allow_html=True
#     )


#     st.markdown('<div class="main-title">SMARTFILL AI</div>', unsafe_allow_html=True)
#     st.divider()

#     # Centered layout using columns
#     col1, col2, col3 = st.columns([1, 2, 1])
#     with col2:
#         # st.markdown('<div class="login-container">', unsafe_allow_html=True)
#         st.markdown('<div class="login-header">🔐 Login</div>', unsafe_allow_html=True)

#         username = st.text_input("Username", placeholder="Enter your username")
#         password = st.text_input("Password", type="password", placeholder="Enter your password")
#         login_button = st.button("Login")

#         if login_button:
#             if not username or not password:
#                 st.warning("⚠️ Please enter both username and password.")
#                 return

#             try:
#                 if USERS.get(username) == password:
#                     st.session_state.logged_in = True
#                     st.success("✅ Login successful!")
#                     st.rerun()
#                 else:
#                     st.error("❌ Invalid credentials.")
#             except requests.exceptions.RequestException:
#                 st.error("⚠️ Failed to connect to authentication server.")
#         st.markdown('</div>', unsafe_allow_html=True)

# # If not logged in, show login page
# if 'logged_in' not in st.session_state:
#     st.session_state.logged_in = False

# if not st.session_state.logged_in:
#     show_login_page()
#     st.stop()


# # Set page configuration
# st.set_page_config(page_title="SMARTFILL AI", layout="wide")

# # Add custom CSS for styling
# st.markdown("""
# <style>
# /* This overrides the default text with our custom text */
# .st-emotion-cache-1v04i6y p:contains("Drag and drop files here"),
# div.st-emotion-cache-1eqt3ip p:first-child,
# div[data-testid="stFileUploader"] p:first-child,
# div[data-baseweb="file-uploader"] p:first-child,
# section[data-testid="stFileUploader"] p:first-child,
# .uploadedFile p:first-child {
#     font-size: 0;
#     line-height: 0;
# }

# /* Add our text after making the original invisible */
# .st-emotion-cache-1v04i6y p:contains("Drag and drop files here")::before,
# div.st-emotion-cache-1eqt3ip p:first-child::before,
# div[data-testid="stFileUploader"] p:first-child::before,
# div[data-baseweb="file-uploader"] p:first-child::before,
# section[data-testid="stFileUploader"] p:first-child::before,
# .uploadedFile p:first-child::before {
#     content: "Please Upload your files here";
#     font-size: 1rem;
#     line-height: 1.6;
#     display: block;
# }

# /* Dashboard styling */
# .metric-card {
#     background-color: white;
#     border-radius: 5px;
#     padding: 15px;
#     box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
#     transition: all 0.3s;
# }

# .metric-card:hover {
#     transform: translateY(-3px);
#     box-shadow: 0 4px 8px rgba(0, 0, 0, 0.15);
# }

# .metric-value {
#     font-size: 32px;
#     font-weight: bold;
#     margin: 10px 0;
#     color: #1E88E5;
# }

# .metric-label {
#     font-weight: 500;
#     color: #616161;
#     margin-top: 0;
# }

# .metric-trend {
#     font-size: 14px;
#     margin-top: 5px;
#     color: #4CAF50;
# }

# .metric-trend-down {
#     color: #F44336;
# }

# .activity-item {
#     padding: 10px 15px;
#     border-left: 4px solid #1E88E5;
#     margin-bottom: 10px;
#     background-color: #F5F7FA;
# }

# .dashboard-section {
#     margin-bottom: 20px;
#     padding: 15px;
#     border-radius: 5px;
#     background-color: #f9f9f9;
# }

# /* Tab styling */
# div[data-testid="stHorizontalBlock"] button[kind="primary"] {
#     background-color: #1E88E5;
#     border-color: #1E88E5;
# }

# div[data-testid="stHorizontalBlock"] button:hover {
#     background-color: #1565C0;
#     border-color: #1565C0;
# }

# div.stTabs [data-baseweb="tab-panel"] {
#     padding-top: 20px;
# }

# </style>
# """, unsafe_allow_html=True)

# # Initialize session state
# if 'rag' not in st.session_state:
#     st.session_state.rag = RAGEngine(layout_aware=True)
# if 'rfp_processor' not in st.session_state:
#     st.session_state.rfp_processor = RFPProcessor()
# if 'chat_history' not in st.session_state:
#     st.session_state.chat_history = []
# if 'indexed' not in st.session_state:
#     st.session_state.indexed = False
# if 'needs_context' not in st.session_state:
#     st.session_state.needs_context = []
# if 'answers' not in st.session_state:
#     st.session_state.answers = {}
# if 'current_question_idx' not in st.session_state:
#     st.session_state.current_question_idx = 0
# if 'questions_mode' not in st.session_state:
#     st.session_state.questions_mode = None
# if 'export_status' not in st.session_state:
#     st.session_state.export_status = {
#         'ready': False,
#         'filename': None,
#         'error': None
#     }
# if 'document_generated' not in st.session_state:
#     st.session_state.document_generated = False
# if 'email_form_visible' not in st.session_state:
#     st.session_state.email_form_visible = False
# if 'email_recipient' not in st.session_state:
#     st.session_state.email_recipient = ""
# if 'current_doc_filename' not in st.session_state:
#     st.session_state.current_doc_filename = None
# if 'email_success' not in st.session_state:
#     st.session_state.email_success = False
# if 'email_success_message' not in st.session_state:
#     st.session_state.email_success_message = ""

# # Add dashboard specific session state
# if 'rfp_history' not in st.session_state:
#     # Mock data for RFP history
#     st.session_state.rfp_history = [
#         {"id": "RFP-2025-001", "title": "Cloud Infrastructure Services", "date": "2025-04-30", "status": "Completed", "questions": 34, "autoAnswered": 28},
#         {"id": "RFP-2025-002", "title": "IT Support Services", "date": "2025-05-01", "status": "Completed", "questions": 42, "autoAnswered": 36},
#         {"id": "RFP-2025-003", "title": "Data Analytics Platform", "date": "2025-05-03", "status": "In Progress", "questions": 27, "autoAnswered": 18},
#         {"id": "RFP-2025-004", "title": "Cybersecurity Assessment", "date": "2025-05-05", "status": "In Progress", "questions": 31, "autoAnswered": 22},
#     ]
# if 'activity_log' not in st.session_state:
#     # Mock data for activity log
#     st.session_state.activity_log = [
#         {"timestamp": "2025-05-06 15:32", "action": "RFP Completed", "details": "Cybersecurity Assessment RFP completed and exported"},
#         {"timestamp": "2025-05-06 14:15", "action": "Document Generated", "details": "Generated response document for Data Analytics Platform RFP"},
#         {"timestamp": "2025-05-06 11:24", "action": "RFP Processed", "details": "Processed new RFP for Network Security Evaluation"},
#         {"timestamp": "2025-05-06 09:45", "action": "Training Data Added", "details": "8 new training documents added to knowledge base"}
#     ]
# if 'current_tab' not in st.session_state:
#     st.session_state.current_tab = "Dashboard"

# # Email functions
# def show_email_form():
#     st.session_state.email_form_visible = True
#     st.session_state.email_success = False  # Reset success state when showing form

# def hide_email_form():
#     st.session_state.email_form_visible = False

# def set_email_success(message):
#     st.session_state.email_success = True
#     st.session_state.email_success_message = message
#     time.sleep(3)
#     st.session_state.email_form_visible = False

# def send_email_now(recipient, filename):
#     if recipient.strip():
#         try:
#             from backend.utils import send_email
#             success, message = send_email(
#                 recipient_email=recipient,
#                 subject="SMARTFILL AI - RFP Responses",
#                 body="Attached are the SMARTFILL AI generated RFP responses.",
#                 attachment_path=filename
#             )
#             if success:
#                 success_msg = f"✅ Document successfully delivered to {recipient}!"
#                 set_email_success(success_msg)
#             return success, message
#         except Exception as e:
#             return False, str(e)
#     else:
#         return False, "Please enter a valid email address"

# # Mode selection callbacks
# def set_all_mode():
#     st.session_state.questions_mode = "all"
#     st.session_state.current_question_idx = 0

# def set_single_mode():
#     st.session_state.questions_mode = "single"
#     st.session_state.current_question_idx = 0

# def clear_input_fields():
#     # Clear any existing context input for the current question
#     current_context_key = f"context_{st.session_state.current_question_idx}"
#     if current_context_key in st.session_state:
#         st.session_state[current_context_key] = ""

#     # Clear any existing file upload for the current question
#     current_file_key = f"file_{st.session_state.current_question_idx}"
#     if current_file_key in st.session_state:
#         st.session_state[current_file_key] = None

# # Document generation function
# def generate_document(answers):
#     """Generate a Word document with RFP responses"""
#     try:
#         # Create Word document
#         doc = Document()
#         doc.add_heading("SMARTFILL AI - RFP Responses", 0)
#         doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
#         doc.add_paragraph(f"Total Questions: {len(answers)}")
#         doc.add_paragraph()

#         # Add each Q&A to document
#         for q, a in answers.items():
#             doc.add_heading("Question", level=2)
#             doc.add_paragraph(q)
#             doc.add_heading("Response", level=2)
#             doc.add_paragraph(a)
#             doc.add_paragraph()

#         # Save to file
#         filename = f"RFP_Responses_{uuid.uuid4().hex[:8]}.docx"
#         doc.save(filename)
#         st.session_state.current_doc_filename = filename

#         # Add to activity log
#         now = datetime.now().strftime("%Y-%m-%d %H:%M")
#         st.session_state.activity_log.insert(0, {
#             "timestamp": now,
#             "action": "Document Generated",
#             "details": f"Response document created: {filename}"
#         })

#         return filename, None
#     except Exception as e:
#         return None, str(e)

# # Dashboard metrics
# def get_dashboard_metrics():
#     metrics = {
#         "total_rfps": len(st.session_state.rfp_history),
#         "completed_rfps": sum(1 for rfp in st.session_state.rfp_history if rfp["status"] == "Completed"),
#         "in_progress": sum(1 for rfp in st.session_state.rfp_history if rfp["status"] == "In Progress"),
#         "auto_answer_rate": round(sum(rfp["autoAnswered"] for rfp in st.session_state.rfp_history) /
#                             sum(rfp["questions"] for rfp in st.session_state.rfp_history) * 100, 1)
#     }
#     return metrics

# # Create sidebar
# with st.sidebar:
#     st.image("https://via.placeholder.com/200x80?text=SMARTFILL+AI", width=200)

#     # Navigation
#     st.sidebar.markdown("## 📊 Navigation")
#     selected_tab = st.sidebar.radio("", ["Dashboard", "Process RFP", "Training Data", "Settings"])
#     st.session_state.current_tab = selected_tab

#     st.sidebar.markdown("---")

#     # Quick stats in sidebar
#     if st.session_state.indexed:
#         st.sidebar.markdown("### 📈 Quick Stats")
#         metrics = get_dashboard_metrics()
#         st.sidebar.markdown(f"**Total RFPs:** {metrics['total_rfps']}")
#         st.sidebar.markdown(f"**Completed:** {metrics['completed_rfps']}")
#         st.sidebar.markdown(f"**Auto-answer Rate:** {metrics['auto_answer_rate']}%")

#     # Help section
#     with st.sidebar.expander("❓ Help & Resources"):
#         st.markdown("- [User Guide](https://example.com)")
#         st.markdown("- [FAQ](https://example.com/faq)")
#         st.markdown("- [Support](mailto:support@example.com)")

# # Display success message if email was sent successfully
# if st.session_state.email_success:
#     st.success(st.session_state.email_success_message)
#     time.sleep(3)
#     st.session_state.email_success = False
#     st.rerun()

# # Dashboard view
# if st.session_state.current_tab == "Dashboard":
#     st.title("🤖 SMARTFILL AI - RFP Assistant")

#     # Metrics section
#     metrics = get_dashboard_metrics()
#     st.markdown("### 📊 RFP Status Overview")

#     col1, col2, col3, col4 = st.columns(4)

#     with col1:
#         st.markdown("""
#         <div class="metric-card">
#             <p class="metric-label">Total RFPs</p>
#             <p class="metric-value">{}</p>
#             <p class="metric-trend">↑ 25% from last month</p>
#         </div>
#         """.format(metrics["total_rfps"]), unsafe_allow_html=True)

#     with col2:
#         st.markdown("""
#         <div class="metric-card">
#             <p class="metric-label">Completed</p>
#             <p class="metric-value">{}</p>
#             <p class="metric-trend">↑ 18% from last month</p>
#         </div>
#         """.format(metrics["completed_rfps"]), unsafe_allow_html=True)

#     with col3:
#         st.markdown("""
#         <div class="metric-card">
#             <p class="metric-label">In Progress</p>
#             <p class="metric-value">{}</p>
#             <p class="metric-trend">↓ 5% from last month</p>
#         </div>
#         """.format(metrics["in_progress"]), unsafe_allow_html=True)

#     with col4:
#         st.markdown("""
#         <div class="metric-card">
#             <p class="metric-label">Auto-Answer Rate</p>
#             <p class="metric-value">{}%</p>
#             <p class="metric-trend">↑ 12% from last month</p>
#         </div>
#         """.format(metrics["auto_answer_rate"]), unsafe_allow_html=True)

#     # Charts and tables section
#     st.markdown("---")
#     chart_col1, chart_col2 = st.columns(2)

#     with chart_col1:
#         st.markdown("### 📈 RFP Performance")

#         # Create dataframe for chart
#         rfp_data = pd.DataFrame(st.session_state.rfp_history)
#         rfp_data['completion_rate'] = rfp_data['autoAnswered'] / rfp_data['questions'] * 100

#         # Plot completion rates
#         fig = px.bar(
#             rfp_data,
#             x='id',
#             y='completion_rate',
#             color='completion_rate',
#             color_continuous_scale='Blues',
#             labels={'completion_rate': 'Auto-completion Rate (%)', 'id': 'RFP ID'},
#             title='Auto-completion Rates by RFP'
#         )

#         st.plotly_chart(fig, use_container_width=True)

#     with chart_col2:
#         st.markdown("### 🔄 Recent Activity")

#         # Activity log
#         for activity in st.session_state.activity_log[:5]:
#             st.markdown(f"""
#             <div class="activity-item">
#                 <strong>{activity['timestamp']}</strong> - {activity['action']}<br>
#                 <small>{activity['details']}</small>
#             </div>
#             """, unsafe_allow_html=True)

#         if len(st.session_state.activity_log) > 5:
#             st.markdown("<div style='text-align:center'><a href='#'>View all activity</a></div>", unsafe_allow_html=True)

#     # RFP history table
#     st.markdown("---")
#     st.markdown("### 📋 RFP History")

#     # Convert to dataframe for table
#     rfp_df = pd.DataFrame(st.session_state.rfp_history)
#     rfp_df = rfp_df.rename(columns={
#         'id': 'RFP ID',
#         'title': 'Title',
#         'date': 'Date',
#         'status': 'Status',
#         'questions': 'Questions',
#         'autoAnswered': 'Auto-Answered'
#     })

#     # Add action column
#     rfp_df['Action'] = [
#         '<a href="#" class="btn btn-primary btn-sm">View</a>' for _ in range(len(rfp_df))
#     ]

#     # Style the Status column
#     def highlight_status(val):
#         if val == 'Completed':
#             return 'background-color: #E8F5E9; color: #2E7D32'
#         elif val == 'In Progress':
#             return 'background-color: #FFF8E1; color: #F57F17'
#         else:
#             return ''

#     # Apply the style
#     styled_df = rfp_df.style.applymap(highlight_status, subset=['Status'])

#     # Display the table
#     st.dataframe(styled_df, use_container_width=True, hide_index=True)

#     # Quick actions section
#     st.markdown("---")
#     st.markdown("### ⚡ Quick Actions")

#     quick_col1, quick_col2, quick_col3 = st.columns(3)

#     with quick_col1:
#         if st.button("🆕 New RFP", use_container_width=True):
#             st.session_state.current_tab = "Process RFP"
#             st.rerun()

#     with quick_col2:
#         if st.button("📊 Generate Report", use_container_width=True):
#             st.info("Generating comprehensive RFP performance report...")

#     with quick_col3:
#         if st.button("📚 Manage Knowledge Base", use_container_width=True):
#             st.session_state.current_tab = "Training Data"
#             st.rerun()

# # Process RFP Tab
# elif st.session_state.current_tab == "Process RFP":
#     st.title("🤖 Process New RFP")

#     if not st.session_state.indexed:
#         st.warning("⚠️ Please upload and process training documents first")

#         training_files = st.file_uploader(
#             "Upload your training documents",
#             type=["pdf", "xlsx", "docx", "csv"],
#             accept_multiple_files=True,
#             help="Upload any PDF, Excel, Word, or CSV files for training"
#         )

#         if st.button("Process Training Documents"):
#             if training_files:
#                 with st.spinner("Processing training documents..."):
#                     # Mock processing
#                     progress_bar = st.progress(0)
#                     for i in range(100):
#                         time.sleep(0.05)
#                         progress_bar.progress(i + 1)

#                     # Add to activity log
#                     now = datetime.now().strftime("%Y-%m-%d %H:%M")
#                     st.session_state.activity_log.insert(0, {
#                         "timestamp": now,
#                         "action": "Training Data Added",
#                         "details": f"{len(training_files)} training documents processed"
#                     })

#                     st.session_state.indexed = True
#                     st.success("✅ Training documents processed successfully!")
#                     st.rerun()
#             else:
#                 st.error("Please upload at least one training document")
#     else:
#         # RFP Upload Area
#         st.markdown("### 📑 Upload RFP Document")

#         rfp_file = st.file_uploader("", type=["xlsx"], key="rfp",
#                                   help="Upload your RFP Excel file here")

#         if rfp_file and st.button("Process RFP"):
#             with st.spinner("Processing RFP..."):
#                 # Mock processing
#                 progress_bar = st.progress(0)
#                 for i in range(100):
#                     time.sleep(0.03)
#                     progress_bar.progress(i + 1)

#                 # Reset state for new RFP
#                 st.session_state.questions_mode = None
#                 st.session_state.current_question_idx = 0
#                 st.session_state.answers = {}
#                 st.session_state.document_generated = False
#                 st.session_state.email_form_visible = False

#                 # Mock RFP data for demo
#                 total_questions = 35
#                 auto_answered = 28
#                 needs_context = total_questions - auto_answered

#                 # Add to activity log
#                 now = datetime.now().strftime("%Y-%m-%d %H:%M")
#                 rfp_id = f"RFP-{now.split()[0].replace('-', '')[-4:]}-{len(st.session_state.rfp_history) + 1:03d}"
#                 st.session_state.activity_log.insert(0, {
#                     "timestamp": now,
#                     "action": "RFP Processed",
#                     "details": f"New RFP {rfp_id} processed with {total_questions} questions"
#                 })

#                 # Add to RFP history
#                 st.session_state.rfp_history.append({
#                     "id": rfp_id,
#                     "title": rfp_file.name.replace(".xlsx", ""),
#                     "date": now.split()[0],
#                     "status": "In Progress",
#                     "questions": total_questions,
#                     "autoAnswered": auto_answered
#                 })

#                 # Mock needs_context for demo
#                 st.session_state.needs_context = [f"Question {i}" for i in range(1, needs_context+1)]

#                 # Mock answers for demo
#                 st.session_state.answers = {
#                     f"Auto-answered Question {i}": f"This is an automatically generated answer for question {i}"
#                     for i in range(1, auto_answered+1)
#                 }

#                 st.success(f"✅ RFP processed! {auto_answered} of {total_questions} questions automatically answered.")

#                 if needs_context > 0:
#                     st.warning(f"⚠️ {needs_context} questions need additional context")
#                     st.info("Please select how you would like to provide context for the questions:")
#                 else:
#                     st.session_state.show_final_button = True

#         # Show Get Final Responses button if appropriate
#         if 'show_final_button' in st.session_state and st.session_state.show_final_button:
#             if st.button("Get Final Responses", key="get_final_responses_auto"):
#                 try:
#                     with st.spinner("Generating final document..."):
#                         # Mock document generation
#                         time.sleep(2)
#                         filename = f"RFP_Responses_{uuid.uuid4().hex[:8]}.docx"
#                         st.session_state.current_doc_filename = filename

#                         # Add to activity log
#                         now = datetime.now().strftime("%Y-%m-%d %H:%M")
#                         st.session_state.activity_log.insert(0, {
#                             "timestamp": now,
#                             "action": "Document Generated",
#                             "details": f"Response document created: {filename}"
#                         })

#                         # Update RFP status
#                         if st.session_state.rfp_history:
#                             st.session_state.rfp_history[0]["status"] = "Completed"

#                         st.success("✅ Completed final question and responses generation!")

#                         # Create two columns for export and email options
#                         col1, col2 = st.columns(2)

#                         with col1:
#                             st.download_button(
#                                 label="📥 Export Responses as DOCX",
#                                 data=b"Sample document content",
#                                 file_name=filename,
#                                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                             )

#                         # The email button with on_click callback
#                         with col2:
#                             st.button("✉️ Send via Email", key="email_button_1", on_click=show_email_form)
#                 except Exception as e:
#                     st.error(f"❌ Failed to generate export: {str(e)}")

#         # Question display mode selection
#         if st.session_state.needs_context:
#             col1, col2 = st.columns(2)
#             with col1:
#                 st.checkbox("Show all questions at once", key="show_all",
#                            on_change=set_all_mode,
#                            value=(st.session_state.questions_mode == "all"))
#             with col2:
#                 st.checkbox("Show questions one by one", key="show_one",
#                            on_change=set_single_mode,
#                            value=(st.session_state.questions_mode == "single"))

#             # Handle mode selection conflicts
#             if st.session_state.get("show_all") and st.session_state.get("show_one"):
#                 st.error("Please select only one display mode")
#                 st.session_state.questions_mode = None

#             # Display questions based on selected mode
#             if st.session_state.questions_mode:
#                 unanswered = [q for q in st.session_state.needs_context if q not in st.session_state.answers]

#                 if unanswered:
#                     if st.session_state.questions_mode == "all":
#                         st.markdown("### Questions Needing Additional Context")
#                         question_containers = {}  # Store containers for each question

#                         for question in unanswered:
#                             question_containers[question] = st.container()
#                             with question_containers[question]:
#                                 st.markdown(f"**Q:** {question}")
#                                 col1, col2, col3 = st.columns([2, 1, 1])
#                                 with col1:
#                                     context = st.text_area("Provide additional context:", key=f"context_{question}")
#                                 with col2:
#                                     context_file = st.file_uploader(
#                                         "Or upload a file with context",
#                                         key=f"file_{question}",
#                                         label_visibility="collapsed"
#                                     )
#                                 with col3:
#                                     # Add Update Context button for each question
#                                     if st.button("Update Context", key=f"update_{question}"):
#                                         if context.strip() or context_file:
#                                             st.session_state.answers[question] = f"This is a generated answer for {question} based on the provided context."
#                                             st.success("Answer updated successfully!")

#                                             # Show preview of the answer
#                                             st.markdown("**Generated Answer:**")
#                                             st.markdown(st.session_state.answers[question])

#                     elif st.session_state.questions_mode == "single":
#                         if st.session_state.current_question_idx < len(unanswered):
#                             question = unanswered[st.session_state.current_question_idx]
#                             st.markdown("### Current Question")
#                             st.markdown(f"**Q:** {question}")

#                             # Add question counter
#                             st.markdown(f"Question {st.session_state.current_question_idx + 1} of {len(unanswered)}")

#                             # Generate unique keys for each question's input fields
#                             context_key = f"context_{st.session_state.current_question_idx}"
#                             file_key = f"file_{st.session_state.current_question_idx}"

#                             col1, col2, col3 = st.columns([2, 1, 1])
#                             with col1:
#                                 context = st.text_area("Provide additional context:", key=context_key)
#                             with col2:
#                                 context_file = st.file_uploader(
#                                     "Or upload a file with context",
#                                     key=file_key,
#                                     label_visibility="collapsed"
#                                 )
#                             with col3:
#                                 # Add Update Context button
#                                 if st.button("Update Context"):
#                                     if context.strip() or context_file:
#                                         st.session_state.answers[question] = f"This is a generated answer for {question} based on the provided context."
#                                         st.success("Answer updated successfully!")

#                                         # Show preview of the answer
#                                         st.markdown("**Generated Answer:**")
#                                         st.markdown(st.session_state.answers[question])

#                             # Add navigation buttons in a new row
#                             nav_col1, nav_col2, nav_col3 = st.columns([1, 1, 1])
#                             with nav_col1:
#                                 if st.session_state.current_question_idx > 0:
#                                     if st.button("⬅️ Previous Question"):
#                                         clear_input_fields()  # Clear current input fields
#                                         st.session_state.current_question_idx -= 1
#                                         st.rerun()
#                             with nav_col3:
#                                 if st.session_state.current_question_idx < len(unanswered) - 1:
#                                     if st.button("Next Question ➡️"):
#                                         clear_input_fields()  # Clear current input fields
#                                         st.session_state.current_question_idx += 1
#                                         st.rerun()

#                             # Only show the current question's answer if it exists
#                             if question in st.session_state.answers:
#                                 st.markdown("**Current Answer:**")
#                                 st.markdown(st.session_state.answers[question])
#                         else:
#                             st.success("✅ All questions have been processed!")
#                 else:
#                     st.success("✅ All questions have been processed!")

#                 # Show Get Final Responses button only when appropriate
#                 if not st.session_state.needs_context or not unanswered:
#                     st.markdown("---")
#                     if st.button("Get Final Responses", key="get_final_responses"):
#                         try:
#                             with st.spinner("Generating final document..."):
#                                 # Mock document generation
#                                 time.sleep(2)
#                                 filename = f"RFP_Responses_{uuid.uuid4().hex[:8]}.docx"
#                                 st.session_state.current_doc_filename = filename

#                                 # Add to activity log
#                                 now = datetime.now().strftime("%Y-%m-%d %H:%M")
#                                 st.session_state.activity_log.insert(0, {
#                                     "timestamp": now,
#                                     "action": "Document Generated",
#                                     "details": f"Response document created: {filename}"
#                                 })

#                                 # Update RFP status
#                                 if st.session_state.rfp_history:
#                                     st.session_state.rfp_history[0]["status"] = "Completed"

#                                 st.success("✅ Completed final question and responses generation!")

#                                 # Create two columns for export and email options
#                                 col1, col2 = st.columns(2)

#                                 with col1:
#                                     st.download_button(
#                                         label="📥 Export Responses as DOCX",
#                                         data=b"Sample document content",
#                                         file_name=filename,
#                                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                                     )

#                                 # The email button uses the on_click parameter for the callback
#                                 with col2:
#                                     st.button("✉️ Send via Email", key="email_button_2", on_click=show_email_form)
#                         except Exception as e:
#                             st.error(f"❌ Failed to generate export: {str(e)}")

# # Training Data Tab
# elif st.session_state.current_tab == "Training Data":
#     st.title("📚 Training Data Management")

#     # Training data statistics
#     st.markdown("### 📊 Knowledge Base Statistics")

#     # Mock data for knowledge base stats
#     kb_stats = {
#         "total_documents": 147,
#         "total_chunks": 2348,
#         "document_types": {
#             "PDF": 53,
#             "DOCX": 42,
#             "XLSX": 28,
#             "CSV": 24
#         },
#         "last_updated": "2025-05-06 09:45"
#     }

#     kb_col1, kb_col2, kb_col3, kb_col4 = st.columns(4)

#     with kb_col1:
#         st.markdown("""
#         <div class="metric-card">
#             <p class="metric-label">Total Documents</p>
#             <p class="metric-value">{}</p>
#         </div>
#         """.format(kb_stats["total_documents"]), unsafe_allow_html=True)

#     with kb_col2:
#         st.markdown("""
#         <div class="metric-card">
#             <p class="metric-label">Total Chunks</p>
#             <p class="metric-value">{}</p>
#         </div>
#         """.format(kb_stats["total_chunks"]), unsafe_allow_html=True)

#     with kb_col3:
#         st.markdown("""
#         <div class="metric-card">
#             <p class="metric-label">Document Types</p>
#             <p class="metric-value">{}</p>
#         </div>
#         """.format(len(kb_stats["document_types"])), unsafe_allow_html=True)

#     with kb_col4:
#         st.markdown("""
#         <div class="metric-card">
#             <p class="metric-label">Last Updated</p>
#             <p class="metric-value" style="font-size: 18px;">{}</p>
#         </div>
#         """.format(kb_stats["last_updated"]), unsafe_allow_html=True)

#     # Document breakdown chart
#     st.markdown("### 📈 Document Type Breakdown")

#     doc_types = list(kb_stats["document_types"].keys())
#     doc_counts = list(kb_stats["document_types"].values())

#     fig = px.pie(
#         values=doc_counts,
#         names=doc_types,
#         title="Document Types Distribution",
#         color_discrete_sequence=px.colors.sequential.Blues[::2],
#         hole=0.4
#     )

#     fig.update_traces(textinfo='percent+label')

#     st.plotly_chart(fig, use_container_width=True)

#     # Upload new training documents
#     st.markdown("### 📤 Upload New Training Documents")

#     training_files = st.file_uploader(
#         "Upload documents to train the system",
#         type=["pdf", "xlsx", "docx", "csv"],
#         accept_multiple_files=True,
#         help="Upload any PDF, Excel, Word, or CSV files for training"
#     )

#     if st.button("Process Documents"):
#         if training_files:
#             with st.spinner("Processing documents..."):
#                 # Mock processing
#                 progress_bar = st.progress(0)
#                 for i in range(100):
#                     time.sleep(0.05)
#                     progress_bar.progress(i + 1)

#                 # Add to activity log
#                 now = datetime.now().strftime("%Y-%m-%d %H:%M")
#                 st.session_state.activity_log.insert(0, {
#                     "timestamp": now,
#                     "action": "Training Data Added",
#                     "details": f"{len(training_files)} training documents processed"
#                 })

#                 st.session_state.indexed = True
#                 st.success(f"✅ Successfully processed {len(training_files)} training documents!")

#                 # Update mock stats
#                 kb_stats["total_documents"] += len(training_files)
#                 kb_stats["total_chunks"] += len(training_files) * 15  # Assume 15 chunks per document
#                 kb_stats["last_updated"] = now

#                 for file in training_files:
#                     ext = file.name.split('.')[-1].upper()
#                     if ext in kb_stats["document_types"]:
#                         kb_stats["document_types"][ext] += 1
#                     else:
#                         kb_stats["document_types"][ext] = 1
#         else:
#             st.error("Please upload at least one document")

#     # Manage existing documents
#     st.markdown("### 🗂️ Manage Existing Documents")

#     # Mock data for document list
#     documents = [
#         {"name": "Product-Specifications-2025.pdf", "type": "PDF", "date_added": "2025-04-12", "size": "2.4 MB"},
#         {"name": "Technical-Architecture.docx", "type": "DOCX", "date_added": "2025-04-15", "size": "1.8 MB"},
#         {"name": "Pricing-Matrix-2025.xlsx", "type": "XLSX", "date_added": "2025-04-22", "size": "1.2 MB"},
#         {"name": "Security-Compliance.pdf", "type": "PDF", "date_added": "2025-04-24", "size": "3.5 MB"},
#         {"name": "Customer-References.docx", "type": "DOCX", "date_added": "2025-04-28", "size": "0.9 MB"},
#         {"name": "Service-Level-Agreements.pdf", "type": "PDF", "date_added": "2025-05-01", "size": "1.7 MB"},
#         {"name": "Quality-Metrics.csv", "type": "CSV", "date_added": "2025-05-03", "size": "0.6 MB"},
#         {"name": "Implementation-Timeline.xlsx", "type": "XLSX", "date_added": "2025-05-05", "size": "1.1 MB"}
#     ]

#     # Convert to dataframe for table display
#     doc_df = pd.DataFrame(documents)

#     # Add search and filter
#     search_col1, search_col2 = st.columns([3, 1])

#     with search_col1:
#         search_term = st.text_input("🔍 Search documents", placeholder="Enter keywords...")

#     with search_col2:
#         filter_type = st.selectbox("Filter by type", ["All"] + list(set(doc_df["type"])))

#     # Apply filters (mock implementation)
#     filtered_df = doc_df
#     if search_term:
#         filtered_df = filtered_df[filtered_df["name"].str.contains(search_term, case=False)]

#     if filter_type != "All":
#         filtered_df = filtered_df[filtered_df["type"] == filter_type]

#     # Display table
#     st.dataframe(filtered_df, use_container_width=True, hide_index=True)

#     # Action buttons
#     doc_action_col1, doc_action_col2, doc_action_col3 = st.columns(3)

#     with doc_action_col1:
#         if st.button("🗑️ Delete Selected", use_container_width=True):
#             st.warning("This will permanently remove the selected documents. Are you sure?")
#             confirm_col1, confirm_col2 = st.columns(2)
#             with confirm_col1:
#                 if st.button("✅ Yes, delete"):
#                     st.success("Selected documents deleted successfully")
#             with confirm_col2:
#                 if st.button("❌ Cancel"):
#                     st.info("Deletion cancelled")

#     with doc_action_col2:
#         if st.button("📥 Export List", use_container_width=True):
#             st.info("Exporting document list...")

#     with doc_action_col3:
#         if st.button("🔄 Refresh Knowledge Base", use_container_width=True):
#             with st.spinner("Refreshing knowledge base..."):
#                 # Mock refresh
#                 progress_bar = st.progress(0)
#                 for i in range(100):
#                     time.sleep(0.03)
#                     progress_bar.progress(i + 1)
#                 st.success("Knowledge base refreshed successfully!")

# # Settings Tab
# elif st.session_state.current_tab == "Settings":
#     st.title("⚙️ Settings")

#     # System settings
#     st.markdown("### 🖥️ System Settings")

#     system_settings = {
#         "model": "gpt-4",
#         "chunk_size": 1024,
#         "chunk_overlap": 200,
#         "temperature": 0.7,
#         "max_tokens": 8192
#     }

#     model_col1, model_col2 = st.columns(2)

#     with model_col1:
#         model = st.selectbox("AI Model", ["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo"])

#     with model_col2:
#         temperature = st.slider("Temperature", min_value=0.0, max_value=1.0, value=system_settings["temperature"], step=0.1)

#     chunking_col1, chunking_col2 = st.columns(2)

#     with chunking_col1:
#         chunk_size = st.number_input("Chunk Size", min_value=256, max_value=2048, value=system_settings["chunk_size"], step=128)

#     with chunking_col2:
#         chunk_overlap = st.number_input("Chunk Overlap", min_value=0, max_value=512, value=system_settings["chunk_overlap"], step=32)

#     max_tokens = st.slider("Max Response Tokens", min_value=1024, max_value=16384, value=system_settings["max_tokens"], step=1024)

#     if st.button("Save System Settings"):
#         st.success("System settings saved successfully!")

#     # Email Settings
#     st.markdown("---")
#     st.markdown("### ✉️ Email Settings")

#     email_settings = {
#         "sender_email": "rfp@smartfill.ai",
#         "smtp_server": "smtp.example.com",
#         "smtp_port": 587,
#         "use_tls": True
#     }

#     email_col1, email_col2 = st.columns(2)

#     with email_col1:
#         sender_email = st.text_input("Sender Email", value=email_settings["sender_email"])

#     with email_col2:
#         reply_to = st.text_input("Reply-To Email", value="support@smartfill.ai")

#     smtp_col1, smtp_col2, smtp_col3 = st.columns(3)

#     with smtp_col1:
#         smtp_server = st.text_input("SMTP Server", value=email_settings["smtp_server"])

#     with smtp_col2:
#         smtp_port = st.number_input("SMTP Port", min_value=1, max_value=65535, value=email_settings["smtp_port"])

#     with smtp_col3:
#         use_tls = st.checkbox("Use TLS", value=email_settings["use_tls"])

#     smtp_user = st.text_input("SMTP Username")
#     smtp_pass = st.text_input("SMTP Password", type="password")

#     if st.button("Save Email Settings"):
#         st.success("Email settings saved successfully!")

#     # Test email
#     st.markdown("---")
#     test_email = st.text_input("Test Email Address")

#     if st.button("Send Test Email") and test_email:
#         with st.spinner("Sending test email..."):
#             time.sleep(2)
#             st.success(f"Test email sent successfully to {test_email}")

#     # User management
#     st.markdown("---")
#     st.markdown("### 👥 User Management")

#     # Mock user data
#     users = [
#         {"name": "John Smith", "email": "john@example.com", "role": "Admin", "last_login": "2025-05-06 14:32"},
#         {"name": "Sarah Johnson", "email": "sarah@example.com", "role": "Editor", "last_login": "2025-05-05 09:15"},
#         {"name": "Michael Brown", "email": "michael@example.com", "role": "Viewer", "last_login": "2025-05-02 11:43"}
#     ]

#     user_df = pd.DataFrame(users)
#     st.dataframe(user_df, use_container_width=True, hide_index=True)

#     user_col1, user_col2 = st.columns(2)

#     with user_col1:
#         if st.button("➕ Add User", use_container_width=True):
#             st.info("Opening user creation form...")

#     with user_col2:
#         if st.button("🔄 Manage Permissions", use_container_width=True):
#             st.info("Opening permissions management...")

# # EMAIL FORM SECTION - This is a completely separate section at the end
# # Only show if the email form should be visible
# if st.session_state.email_form_visible and st.session_state.current_doc_filename:
#     # Create a container with a border to make it stand out
#     email_container = st.container()

#     with email_container:
#         st.markdown("""
#         <style>
#         .email-section {
#             border: 1px solid #ddd;
#             border-radius: 5px;
#             padding: 15px;
#             margin: 10px 0;
#             background-color: #f8f9fa;
#         }
#         .success-message {
#             background-color: #d4edda;
#             color: #155724;
#             padding: 10px;
#             border-radius: 5px;
#             margin: 10px 0;
#             text-align: center;
#             font-weight: bold;
#             animation: fadeIn 0.5s;
#         }
#         @keyframes fadeIn {
#             from { opacity: 0; }
#             to { opacity: 1; }
#         }
#         </style>
#         <div class="email-section">
#         <h3>✉️ Send Document via Email</h3>
#         </div>
#         """, unsafe_allow_html=True)

#         # Email input field
#         recipient_email = st.text_input("Enter recipient's email address:", key="email_recipient_input")

#         # Create columns for send and cancel buttons
#         button_col1, button_col2, button_col3 = st.columns([1, 1, 2])

#         with button_col1:
#             if st.button("📨 Send Email", key="send_email_btn"):
#                 if recipient_email:
#                     with st.spinner("Sending email..."):
#                         # Mock sending process for demo
#                         time.sleep(2)
#                         success = True

#                         if success:
#                             # Display a beautiful success message
#                             st.markdown(f"""
#                             <div class="success-message">
#                                 <div>✅ Document successfully delivered to {recipient_email}!</div>
#                                 <div style="font-size: 0.9em; margin-top: 5px;">The recipient will receive the document shortly.</div>
#                             </div>
#                             """, unsafe_allow_html=True)

#                             # Set the success state and schedule the form to hide
#                             st.session_state.email_success = True
#                             st.session_state.email_success_message = f"✅ Document successfully delivered to {recipient_email}!"
#                             st.session_state.email_form_visible = False

#                             # Add to activity log
#                             now = datetime.now().strftime("%Y-%m-%d %H:%M")
#                             st.session_state.activity_log.insert(0, {
#                                 "timestamp": now,
#                                 "action": "Email Sent",
#                                 "details": f"Document sent to {recipient_email}"
#                             })

#                             st.rerun()
#                         else:
#                             st.error(f"❌ Failed to send email")
#                 else:
#                     st.warning("⚠️ Please enter a recipient email address")

#         with button_col2:
#             if st.button("❌ Cancel", key="cancel_email_btn"):
#                 st.session_state.email_form_visible = False
#                 st.rerun()
