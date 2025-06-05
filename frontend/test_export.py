import streamlit as st
import time
import os
from docx import Document
import uuid

st.set_page_config(page_title="SMARTFILL AI - Export Test", layout="wide")

# 1. Simulate some answers (mock GPT output)
if "answers" not in st.session_state:
    st.session_state.answers = {
        "What is your company's name?": "Kaizen Health",
        "Do you operate in Texas?": "Yes, we serve several counties in Texas.",
        "Do you offer wheelchair-accessible vehicles?": "Yes, ADA-compliant vehicles are part of our fleet.",
    }

# 2. Show answers on screen
st.subheader("🧠 Generated Answers Preview")
for q, a in st.session_state.answers.items():
    st.markdown(f"**Q:** {q}")
    st.markdown(f"**A:** {a}")
    st.markdown("---")

# 3. Trigger export
if st.button("📝 Generate Final Responses"):
    with st.spinner("Generating final responses..."):
        try:
            # Create a new Word document
            doc = Document()
            doc.add_heading("SMARTFILL AI RFP Responses", 0)

            doc.add_heading("Response Details", level=1)
            doc.add_paragraph(f"Generated Date: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Total Questions: {len(st.session_state.answers)}")
            doc.add_paragraph()

            for q, a in st.session_state.answers.items():
                doc.add_paragraph(f"Question: {q}", style="List Bullet")
                doc.add_paragraph(f"Answer: {a}")
                doc.add_paragraph()

            # Save the document
            output_filename = (
                f"RFP_Responses_{time.strftime('%Y%m%d')}_{uuid.uuid4().hex[:8]}.docx"
            )
            doc.save(output_filename)

            # Show download button if file was created
            if os.path.exists(output_filename):
                st.success("✅ Document generated successfully!")
                with open(output_filename, "rb") as f:
                    st.download_button(
                        label="📥 Download RFP Response Document",
                        data=f,
                        file_name=output_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            else:
                st.error("❌ Document file not found after saving.")

        except Exception as e:
            st.error(f"❌ An error occurred: {str(e)}")
