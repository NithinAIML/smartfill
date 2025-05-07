import faiss
import numpy as np
import pickle
from openai import OpenAI
import os
from dotenv import load_dotenv
from docx import Document
from openpyxl import load_workbook

# Load environment variables
load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def get_embedding(text):
    """Get embedding for a text using OpenAI's API"""
    response = client.embeddings.create(
        model="text-embedding-ada-002",
        input=text
    )
    return response.data[0].embedding

def load_faiss_and_chunks():
    """Load FAISS index and document chunks"""
    print("Loading FAISS index and chunks...")
    index = faiss.read_index("faiss_index.index")
    with open("doc_chunks.pkl", "rb") as f:
        chunks = pickle.load(f)
    
    # Debug information
    print(f"FAISS index dimension: {index.d}")
    print(f"Number of vectors in index: {index.ntotal}")
    print(f"Number of chunks: {len(chunks)}")
    
    # Verify index is not empty
    if index.ntotal == 0:
        raise ValueError("FAISS index is empty")
    
    return index, chunks

def query_similar_chunks(question, index, chunks, k=3):
    """Query FAISS index to get most similar chunks"""
    question_embedding = np.array([get_embedding(question)]).astype('float32')
    
    # Debug dimensionality
    print(f"\nDebug - Query embedding shape: {question_embedding.shape}")
    print(f"Debug - FAISS index dimension: {index.d}")
    
    # Verify embedding dimension matches index
    if question_embedding.shape[1] != index.d:
        raise ValueError(f"Embedding dimension {question_embedding.shape[1]} does not match index dimension {index.d}")
    
    D, I = index.search(question_embedding, k)
    
    # Debug information
    print(f"Debug - Distances: {D}")
    print(f"Debug - Indices: {I}")
    
    # Check if any matches were found
    if len(I[0]) == 0:
        print(f"Warning: No similar chunks found for question: {question}")
        return []
    
    # Get chunks and filter out any None values
    similar_chunks = [chunks[i] for i in I[0] if i < len(chunks)]
    if not similar_chunks:
        print(f"Warning: No valid chunks found for question: {question}")
        return []
        
    return similar_chunks

def generate_answer(question, context_chunks):
    """Generate answer using GPT-4 with retrieved context"""
    context = "\n\n".join(context_chunks)
    prompt = f"""Based on the following context, answer the question.
    
Context:
{context}

Question: {question}

Answer:"""
    
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a knowledgeable assistant providing accurate answers based on given context."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7,
        max_tokens=500
    )
    
    return response.choices[0].message.content

def main():
    # Load FAISS index and chunks
    index, chunks = load_faiss_and_chunks()
    
    # Load questions from Excel
    wb = load_workbook("RFP-Document.xlsx")
    ws = wb["RFP Questions"]
    
    questions = []
    for row in ws.iter_rows(min_row=0):
        cell = row[1]
        if cell.value and not cell.font.bold:
            questions.append(cell.value.strip())
    
    # Process each question
    answers = {}
    for question in questions:
        try:
            # Get relevant context chunks
            similar_chunks = query_similar_chunks(question, index, chunks)
            
            # Skip if no chunks found
            if not similar_chunks:
                print(f"Skipping question due to no relevant context: {question}")
                answers[question] = "Unable to generate answer - no relevant context found"
                continue
            
            # Generate answer
            answer = generate_answer(question, similar_chunks)
            answers[question] = answer
            
            # Print progress
            print(f"Q: {question}")
            print(f"A: {answer}")
            print("-" * 80)
            
        except Exception as e:
            print(f"Error processing question '{question}': {str(e)}")
            answers[question] = f"Error: {str(e)}"
    
    # Save results to a Word document
    doc = Document()
    doc.add_heading('RFP Response Document', 0)
    
    for question, answer in answers.items():
        doc.add_heading(question, level=1)
        doc.add_paragraph(answer)
        doc.add_paragraph()  # Add spacing
    
    doc.save('RFP_Responses.docx')
    print("Responses saved to 'RFP_Responses.docx'")

if __name__ == "__main__":
    main()