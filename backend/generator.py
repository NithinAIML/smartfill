import os
import uuid
from docx import Document

OUTPUT_DIR = "generated_outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

async def generate_output_doc(questions, answers):
    doc = Document()
    doc.add_heading("RFP Response Document", 0)

    for i, question in enumerate(questions, start=1):
        doc.add_heading(f"Q{i}. {question}", level=2)
        answer = answers.get(question, "No answer available.")
        doc.add_paragraph(answer)

    file_name = f"rfp_response_{uuid.uuid4().hex[:8]}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)
    doc.save(file_path)

    return file_path
